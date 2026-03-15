"""
Word export service for OpenClaw review tasks.

Generates a temporary .docx export for optional external editing.
Word is NOT a source of truth — all review decisions must be made
through the application UI.

Architecture note:
    build_document() is a pure function returning a Document object.
    export_and_open() handles I/O side effects.
    This separation makes future import-from-Word integration straightforward
    without modifying the core build logic.
"""
from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional

import structlog

log = structlog.get_logger(__name__)

EXPORTS_TMP_DIR = Path("exports") / "tmp"


def get_export_path(task_id: int) -> Path:
    """Return a safe, deterministic file path for a task export."""
    EXPORTS_TMP_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORTS_TMP_DIR / f"review_task_{task_id}.docx"


def build_document(
    task_id: int,
    review_status: str,
    chart_entry: Optional[Dict[str, Any]] = None,
    youtube_video: Optional[Dict[str, Any]] = None,
    final_artist: str = "",
    final_title: str = "",
    final_lyrics_text: str = "",
    review_notes: str = "",
) -> Any:  # docx.Document
    """
    Build a python-docx Document for the given review task data.
    Pure function — does not write to disk.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Header
    doc.add_heading(f"OpenClaw Review Task #{task_id}", level=1)
    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run(review_status)
    p.add_run("   |   Exported: ").bold = False
    p.add_run(datetime.now(tz=timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    doc.add_paragraph()

    # Chart Source
    doc.add_heading("Chart Source", level=2)
    ce = chart_entry or {}
    _kv(doc, "Position", str(ce.get("chart_position") or "—"))
    _kv(doc, "Artist (raw)", ce.get("artist_raw") or "—")
    _kv(doc, "Title (raw)", ce.get("song_title_raw") or "—")
    _kv(doc, "YouTube URL", ce.get("youtube_url") or "—")
    doc.add_paragraph()

    # YouTube Metadata
    doc.add_heading("YouTube Metadata", level=2)
    yt = youtube_video or {}
    _kv(doc, "Video ID", yt.get("youtube_video_id") or "—")
    _kv(doc, "Video Title", yt.get("video_title") or "—")
    _kv(doc, "Channel", yt.get("channel_title") or "—")
    _kv(doc, "Published", (yt.get("published_at") or "")[:10] or "—")
    _kv(doc, "Canonical URL", yt.get("canonical_url") or "—")
    doc.add_paragraph()

    # Raw Description
    doc.add_heading("Raw Description (read-only reference)", level=2)
    desc = yt.get("description_raw") or ""
    doc.add_paragraph(desc if desc else "(no description)")
    doc.add_paragraph()

    # Final Fields
    doc.add_heading("Final Fields (approve only in the app)", level=2)
    _kv(doc, "Final Artist", final_artist or "—")
    _kv(doc, "Final Title", final_title or "—")
    doc.add_paragraph()
    doc.add_paragraph().add_run("Final Lyrics Text:").bold = True
    doc.add_paragraph(final_lyrics_text if final_lyrics_text else "(not yet entered)")
    if review_notes:
        doc.add_paragraph()
        _kv(doc, "Review Notes", review_notes)

    # Footer notice
    doc.add_paragraph()
    notice_para = doc.add_paragraph(
        "⚠  This document is an export only. "
        "Approve / Reject must be done in the OpenClaw application. "
        "Word is NOT a source of truth. Import from Word is not supported."
    )
    notice_para.runs[0].italic = True

    return doc


def export_and_open(
    task_id: int,
    review_status: str,
    chart_entry: Optional[Dict[str, Any]] = None,
    youtube_video: Optional[Dict[str, Any]] = None,
    final_artist: str = "",
    final_title: str = "",
    final_lyrics_text: str = "",
    review_notes: str = "",
) -> str:
    """
    Build, save, and open a .docx export via the Windows shell.
    Returns the path to the saved file.
    Raises RuntimeError on failure (save or open).
    """
    doc = build_document(
        task_id=task_id,
        review_status=review_status,
        chart_entry=chart_entry,
        youtube_video=youtube_video,
        final_artist=final_artist,
        final_title=final_title,
        final_lyrics_text=final_lyrics_text,
        review_notes=review_notes,
    )

    path = get_export_path(task_id)
    try:
        doc.save(str(path))
        log.info("word_export_saved", task_id=task_id, path=str(path))
    except Exception as exc:
        raise RuntimeError(f"Failed to save export file: {exc}") from exc

    try:
        os.startfile(str(path.resolve()))
        log.info("word_export_opened", task_id=task_id, path=str(path))
    except Exception as exc:
        raise RuntimeError(
            f"Export saved to '{path}' but could not be opened: {exc}\n"
            "Open the file manually."
        ) from exc

    return str(path)


def _kv(doc: Any, key: str, value: str) -> None:
    """Add a 'Key: value' line to the document."""
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)
