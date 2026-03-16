"""
Unified Word export service for OpenClaw review pipeline.

Exports multiple reviewed songs into a single .docx file with:
  - Static table of contents (list of songs)
  - One section per song: heading + lyrics

TOC NOTE: python-docx cannot generate auto-updating Word TOC field codes
reliably. We use a static TOC (plain-text list in beginning of document).
The operator can update it manually in Word via References → Update Table
if needed, but the static list is already fully readable.

Word is NOT a source of truth. Export only — no import supported.
"""
from __future__ import annotations

import os
from dataclasses import dataclass, field
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

import structlog

log = structlog.get_logger(__name__)

EXPORTS_COLLECTIONS_DIR = Path("exports") / "collections"

_EXPORTABLE_STATUSES: Set[str] = {"approved", "approved_with_edits"}


# ---------------------------------------------------------------------------
# Data transfer objects
# ---------------------------------------------------------------------------


@dataclass
class ExportItem:
    """All data needed to render one song section."""
    task_id: int
    chart_position: Optional[int]
    final_artist: str
    final_song_title: str
    final_lyrics_text: str
    review_status: str


@dataclass
class ExportPolicy:
    """Controls which items are included and how the document is formatted."""
    included_statuses: Set[str] = field(
        default_factory=lambda: {"approved", "approved_with_edits"}
    )
    skip_empty_lyrics: bool = True
    sort_by_position: bool = True
    page_breaks: bool = True


@dataclass
class ExportSummary:
    path: str
    exported: int
    skipped: int
    skipped_status: int
    skipped_empty: int
    total_input: int


# ---------------------------------------------------------------------------
# Filtering / sorting
# ---------------------------------------------------------------------------


def filter_and_sort(
    items: List[ExportItem],
    policy: ExportPolicy,
) -> Tuple[List[ExportItem], ExportSummary]:
    """
    Apply policy to items. Returns (included_items, summary).
    Summary contains counts; path is filled in later by export_collection().
    """
    included: List[ExportItem] = []
    skipped_status = 0
    skipped_empty = 0

    for item in items:
        if item.review_status not in policy.included_statuses:
            skipped_status += 1
            continue
        if policy.skip_empty_lyrics and not (item.final_lyrics_text or "").strip():
            skipped_empty += 1
            continue
        included.append(item)

    if policy.sort_by_position:
        included.sort(key=lambda i: (i.chart_position if i.chart_position is not None else 9999, i.task_id))

    skipped = skipped_status + skipped_empty
    summary = ExportSummary(
        path="",
        exported=len(included),
        skipped=skipped,
        skipped_status=skipped_status,
        skipped_empty=skipped_empty,
        total_input=len(items),
    )
    return included, summary


# ---------------------------------------------------------------------------
# Document builder (pure — no disk I/O)
# ---------------------------------------------------------------------------


def _section_heading(item: ExportItem) -> str:
    """Format: 'Position 17. Artist - Title'"""
    pos = item.chart_position if item.chart_position is not None else "?"
    artist = item.final_artist or "—"
    title = item.final_song_title or "—"
    return f"Position {pos}. {artist} - {title}"


def _add_bookmark(paragraph: Any, bookmark_id: int, bookmark_name: str) -> None:
    """Wrap paragraph content with a Word bookmark."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))

    p_elem = paragraph._p
    p_elem.insert(0, bm_start)  # before runs
    p_elem.append(bm_end)


def _add_toc_hyperlink(paragraph: Any, text: str, anchor: str) -> None:
    """Add a clickable hyperlink in TOC that jumps to anchor bookmark."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Remove existing runs from paragraph
    p_elem = paragraph._p
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)

    # Create run inside hyperlink
    run_elem = OxmlElement('w:r')

    # Apply hyperlink character style
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run_elem.append(rPr)

    # Add text
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run_elem.append(t_elem)
    hyperlink.append(run_elem)
    p_elem.append(hyperlink)


def _set_rtl(paragraph: Any) -> None:
    """Set paragraph and run direction to RTL for Hebrew content."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    # Also set jc to right
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'right')
    # Set rtl on each run
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)


def build_document(items: List[ExportItem], policy: ExportPolicy) -> Any:  # docx.Document
    """
    Build a python-docx Document from the given items.
    Pure function — no disk I/O.

    Structure:
      1. Document title
      2. Static TOC (clickable hyperlinks to bookmarks)
      3. Per-song sections: heading (with bookmark) + lyrics (RTL)
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # ── Title ───────────────────────────────────────────────────────────────
    doc.add_heading("Top 100 Lyrics Collection", level=1)

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"Exported: {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  |  "
        f"Songs: {len(items)}"
    ).italic = True
    doc.add_paragraph()

    # ── Static TOC with clickable hyperlinks ────────────────────────────────
    doc.add_heading("Table of Contents", level=2)

    for idx, item in enumerate(items):
        p = doc.add_paragraph(style="Normal")
        p.add_run("")  # placeholder run — will be replaced by hyperlink
        heading_text = _section_heading(item)
        anchor = f"song_{idx}"
        _add_toc_hyperlink(p, heading_text, anchor)

    doc.add_paragraph()  # spacer

    # ── Song sections ────────────────────────────────────────────────────────
    for idx, item in enumerate(items):
        # Heading for song (Heading 2 so it shows up in Word's nav panel)
        heading_text = _section_heading(item)
        song_heading = doc.add_heading(heading_text, level=2)

        # Add bookmark to heading for TOC hyperlink targets
        _add_bookmark(song_heading, idx, f"song_{idx}")

        # Apply RTL to heading
        _set_rtl(song_heading)

        # Lyrics — preserve newlines from final_lyrics_text
        lyrics = (item.final_lyrics_text or "").strip()
        if lyrics:
            for line in lyrics.splitlines():
                p = doc.add_paragraph(line if line.strip() else "")
                p.paragraph_format.space_after = Pt(0)
                _set_rtl(p)
        else:
            p = doc.add_paragraph("(no lyrics)")
            p.runs[0].italic = True
            _set_rtl(p)

        # Page break before next song (but not after last)
        if policy.page_breaks and idx < len(items) - 1:
            doc.add_page_break()
        else:
            doc.add_paragraph()  # spacer

    return doc


# ---------------------------------------------------------------------------
# File I/O
# ---------------------------------------------------------------------------


def get_export_path(label: str = "") -> Path:
    """Return a safe, timestamped file path in exports/collections/."""
    EXPORTS_COLLECTIONS_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().isoformat()
    safe_label = label.replace(" ", "_").replace("/", "-")[:30] if label else "export"
    return EXPORTS_COLLECTIONS_DIR / f"{safe_label}_{today}.docx"


def export_collection(
    items: List[ExportItem],
    policy: ExportPolicy,
    label: str = "top100lyrics",
) -> ExportSummary:
    """
    Filter items, build document, save to disk.
    Returns ExportSummary with path filled in.
    Raises ValueError if nothing to export.
    Raises RuntimeError on save failure.
    """
    included, summary = filter_and_sort(items, policy)

    if not included:
        raise ValueError("No eligible items to export after applying policy.")

    doc = build_document(included, policy)
    path = get_export_path(label)

    try:
        doc.save(str(path))
        log.info("unified_word_export_saved", path=str(path), exported=summary.exported)
    except Exception as exc:
        raise RuntimeError(f"Failed to save export: {exc}") from exc

    summary.path = str(path)
    return summary


def open_file(path: str) -> None:
    """Open a file using the Windows shell."""
    os.startfile(os.path.abspath(path))


def open_folder(path: str) -> None:
    """Open the folder containing the file in Windows Explorer."""
    folder = os.path.dirname(os.path.abspath(path))
    os.startfile(folder)
